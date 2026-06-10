import os
import csv
import hashlib
import glob
import pandas as pd
from typing import Generator, List, Dict, Any
from pypdf import PdfReader
from docx import Document
from PIL import Image
import pytesseract
from pdf2image import convert_from_path


# Optional: If Tesseract is not in your system PATH on Windows, uncomment and point to your executable:
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

class DocumentLoader:
    def __init__(self, target_dir="test_files"):
        """
        Initializes the DocumentLoader and ensures the target directory exists.
        """
        self.target_dir = target_dir
        if not os.path.exists(self.target_dir):
            os.makedirs(self.target_dir)

    def _generate_hash(self, content: str) -> str:
        """Generates a unique SHA-256 hash for a given string block."""
        return hashlib.sha256(content.encode('utf-8')).hexdigest()

    def _create_chunk_metadata(self, doc_id: str, file_path: str, content: str, source_type: str, extra_meta: dict) -> dict:
        """Creates a standardized metadata structure for strict traceability."""
        chunk_hash = self._generate_hash(content)
        metadata = {
            "doc_id": doc_id,
            "chunk_id": f"{doc_id}_{chunk_hash[:12]}",
            "file_name": os.path.basename(file_path),
            "source_type": source_type,
            "chunk_hash": chunk_hash
        }
        metadata.update(extra_meta)
        return metadata

    # 1. PDF EXTRACTION (Page Level + Tesseract OCR fallback)
    def load_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Extracts content page-by-page. If a page yields no text, 
        it falls back to Tesseract OCR automatically.
        """
        doc_id = self._generate_hash(file_path)
        chunks = []
        reader = PdfReader(file_path)
        
        for page_idx, page in enumerate(reader.pages):
            page_num = page_idx + 1
            text = page.extract_text() or ""
            text = text.strip()
            
            # Fallback to Tesseract OCR if the page text is empty (likely a scanned PDF)
            if not text:
                try:
                    images = convert_from_path(file_path, first_page=page_num, last_page=page_num)
                    if images:
                        text = pytesseract.image_to_string(images[0]).strip()
                except Exception as e:
                    print(f"⚠️ OCR Fallback failed on {os.path.basename(file_path)} Page {page_num}: {e}")
            
            if text:
                meta = self._create_chunk_metadata(
                    doc_id=doc_id, 
                    file_path=file_path, 
                    content=text, 
                    source_type="pdf", 
                    extra_meta={"page_number": page_num}
                )
                chunks.append({"text": text, "metadata": meta})
                
        return chunks

    # 2. DOCX EXTRACTION
    def load_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Parses a word document element-by-element (paragraphs and tables)
        and packages them into a structural list pattern matching the pipeline layout.
        """
        doc_id = self._generate_hash(file_path)
        chunks = []
        doc = Document(file_path)
        element_idx = 0
        
        for element in doc.element.body:
            if element.tag.endswith('p'):
                p = [p for p in doc.paragraphs if p._element == element]
                if p and p[0].text.strip():
                    text = p[0].text.strip()
                    element_idx += 1
                    meta = self._create_chunk_metadata(
                        doc_id=doc_id, file_path=file_path, content=text, 
                        source_type="docx", extra_meta={"element_index": element_idx, "sub_type": "paragraph"}
                    )
                    chunks.append({"text": text, "metadata": meta})
                    
            elif element.tag.endswith('tbl'):
                t = [t for t in doc.tables if t._element == element]
                if t:
                    element_idx += 1
                    table_data = []
                    for row in t[0].rows:
                        row_data = [cell.text.strip() for cell in row.cells]
                        table_data.append(" | ".join(row_data))
                    
                    text = "\n".join(table_data).strip()
                    if text:
                        meta = self._create_chunk_metadata(
                            doc_id=doc_id, file_path=file_path, content=text, 
                            source_type="docx", extra_meta={"element_index": element_idx, "sub_type": "table"}
                        )
                        chunks.append({"text": text, "metadata": meta})
        return chunks

    # 3. CSV ROW EXTRACTION (using pandas)
    def load_csv(self, file_path: str, format_mode: str = "semantic") -> List[Dict[str, Any]]:
        """
        Parses CSV files and handles every row as an independent tracking chunk using pandas.
        """
        doc_id = self._generate_hash(file_path)
        chunks = []
        file_name = os.path.basename(file_path)
        
        try:
            df = pd.read_csv(file_path)
            for idx, row in df.iterrows():
                row_dict = row.to_dict()
                if format_mode == "semantic":
                    items = [f"'{k}' is '{v}'" for k, v in row_dict.items() if pd.notna(v) and str(v).strip() != ""]
                    if items:
                        row_str = f"In the file '{file_name}', record #{idx + 1} contains: " + ", ".join(items) + "."
                    else:
                        row_str = ""
                else:
                    items = [f"{k}: {v}" for k, v in row_dict.items() if pd.notna(v) and str(v).strip() != ""]
                    row_str = ", ".join(items)
                
                if row_str.strip():
                    meta = self._create_chunk_metadata(
                        doc_id=doc_id, file_path=file_path, content=row_str, 
                        source_type="csv", extra_meta={"row_index": idx + 1}
                    )
                    chunks.append({"text": row_str, "metadata": meta})
        except Exception as e:
            print(f"❌ Failed to parse CSV file {file_name} using pandas: {e}")
            
        return chunks

    # 4. IMAGE TESSERACT OCR LOADING
    def load_image(self, file_path: str) -> List[Dict[str, Any]]:
        """Loads pure image files (PNG/JPG) using Tesseract OCR."""
        doc_id = self._generate_hash(file_path)
        chunks = []
        
        try:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img).strip()
            
            if text:
                meta = self._create_chunk_metadata(
                    doc_id=doc_id, file_path=file_path, content=text, 
                    source_type="image", extra_meta={"type": "scanned_ocr"}
                )
                chunks.append({"text": text, "metadata": meta})
        except Exception as e:
            print(f"❌ Failed to execute OCR processing on image file {os.path.basename(file_path)}: {e}")
            
        return chunks

    # 5. NEW: TXT PASSAGE LOADING
    def load_txt(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Reads plain text files line-by-line or paragraph-by-paragraph, 
        tracking line index numbers for clean citation reference pointers.
        """
        doc_id = self._generate_hash(file_path)
        chunks = []
        
        try:
            with open(file_path, mode='r', encoding='utf-8', errors='ignore') as f:
                lines = f.readlines()
                
            for idx, line in enumerate(lines):
                line_content = line.strip()
                if line_content:  # Skip empty whitespace breaks
                    meta = self._create_chunk_metadata(
                        doc_id=doc_id, 
                        file_path=file_path, 
                        content=line_content, 
                        source_type="txt", 
                        extra_meta={"line_index": idx + 1}
                    )
                    chunks.append({"text": line_content, "metadata": meta})
        except Exception as e:
            print(f"❌ Failed to parse plain text file {os.path.basename(file_path)}: {e}")
            
        return chunks

    # 6. BATCH LOADING WITH CRYPTOGRAPHIC DEDUPLICATION
    def load_batch(self, file_paths: List[str], csv_format_mode: str = "semantic") -> List[Dict[str, Any]]:
        """
        Processes a batch list of multiple file paths. Collects payloads across all 
        parsers and uses a set hash tracker to guarantee exact cryptographic deduplication.
        Handles individual parser failures gracefully to prevent crashing the loop.
        """
        all_chunks = []
        seen_chunk_hashes = set()
        
        for path in file_paths:
            if not os.path.exists(path):
                print(f"⚠️ File path not found, skipping target location: {path}")
                continue
                
            _, ext = os.path.splitext(path)
            ext = ext.lower().replace('.', '')
            file_chunks = []
            
            try:
                if ext == 'pdf':
                    file_chunks = self.load_pdf(path)
                elif ext == 'docx':
                    file_chunks = self.load_docx(path)
                elif ext == 'csv':
                    file_chunks = self.load_csv(path, format_mode=csv_format_mode)
                elif ext == 'txt':
                    file_chunks = self.load_txt(path)  # Integrated TXT mapping route
                elif ext in ['png', 'jpg', 'jpeg']:
                    file_chunks = self.load_image(path)
                else:
                    print(f"⚠️ Unsupported storage system extension format: .{ext}")
                    continue
            except Exception as e:
                print(f"❌ Failed processing file {os.path.basename(path)}: {e}")
                continue
                
            # Process and deduplicate chunks for the file
            for chunk in file_chunks:
                h = chunk["metadata"]["chunk_hash"]
                if h not in seen_chunk_hashes:
                    seen_chunk_hashes.add(h)
                    all_chunks.append(chunk)
                    
        print(f"✅ Ingestion complete. Collected and verified {len(all_chunks)} unique document records.")
        return all_chunks

    # 7. NEW: DIRECTORY AGGREGATION SCAN USING GLOB
    def load_directory(self, folder_path: str, csv_format_mode: str = "semantic") -> List[Dict[str, Any]]:
        """
        Scans a folder using glob and aggregates all supported files.
        """
        pattern = os.path.join(folder_path, "*")
        files = glob.glob(pattern)
        # Exclude directories
        files = [f for f in files if os.path.isfile(f)]
        return self.load_batch(files, csv_format_mode=csv_format_mode)